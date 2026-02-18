from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "PNW_Tender_Proposal_Gambew_Consulting.docx"
PDF_PATH = ROOT / "PNW_Tender_Proposal_Gambew_Consulting.pdf"

TODAY = date(2026, 2, 18)

CONTACT_DETAILS = [
    "Gareth Bew",
    "Gambew Consulting",
    "Email: gambew@outlook.com",
    "Mobile: +27 75 470 2529",
    "WhatsApp: +27 84 0333 777",
]

PUBLIC_INTERFACES = [
    "/incidents",
    "/events",
    "/register",
    "/find",
    "/volunteer",
    "/vacation-watch",
    "/start-scheme",
    "/contact",
]

DATA_DOMAINS = [
    "User",
    "Incident",
    "Event",
    "EventRsvp",
    "Zone",
    "ContactMessage",
    "VolunteerInterest",
    "VacationWatch",
    "SchemeInquiry",
    "SafetyTip",
]

WEEK_PLAN = [
    (
        "Week 1-2: Operational Baseline and Assurance Hardening",
        [
            "Complete platform baseline audit and implementation verification.",
            "Stabilize release process and harden runtime/build reliability.",
            "Close governance gaps in evidence capture and signoff controls.",
        ],
    ),
    (
        "Week 3-4: UX and Content Value Uplift",
        [
            "Refine member and stakeholder journeys for clarity and conversion.",
            "Polish information architecture and executive-facing content quality.",
            "Improve usability, consistency, and service discoverability.",
        ],
    ),
    (
        "Week 5-6: Command Layer and C-Level Visibility",
        [
            "Introduce operational reporting views for committee oversight.",
            "Define KPI framework for service usage, engagement, and quality.",
            "Align data outputs for decision-making cadence and governance.",
        ],
    ),
    (
        "Week 7: Launch Readiness and Stakeholder Signoff",
        [
            "Execute structured verification scenarios across critical routes.",
            "Finalize evidence matrix and executive presentation pack.",
            "Run acceptance walkthrough with committee stakeholders.",
        ],
    ),
    (
        "Week 8: Go-Live, Transition, and Handover Closure",
        [
            "Perform controlled launch and operational handover.",
            "Activate transition support model and response pathways.",
            "Close delivery with formal acceptance and support onboarding.",
        ],
    ),
]

FINANCIAL_TABLE = [
    ["Item", "Commercial Term (ZAR)"],
    [
        "Immediate payment",
        "R50,000 for platform already built, including updates, presentation, verification, and expansion.",
    ],
    ["Completion milestone", "R10,000 upon final acceptance as 'Done'."],
    [
        "Included support period",
        "6 months free support, maintenance, transition assistance, hosting, and technical infrastructure.",
    ],
    [
        "Ongoing service fee",
        "From month 7 onward: R1,500 per month for infrastructure, support, maintenance, and hosting.",
    ],
]


def add_title(doc: Document, text: str, size: int = 24) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def add_subtitle(doc: Document, text: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Normal"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_docx() -> None:
    doc = Document()

    add_title(doc, "GAMBEW CONSULTING")
    add_subtitle(doc, "Professional Tender Submission")
    doc.add_paragraph()
    add_title(doc, "PNW World-Class Platform Transformation Proposal", size=18)
    add_subtitle(doc, "Prepared for: Plumstead Neighbourhood Watch Committee")
    add_subtitle(doc, f"Date: {TODAY.strftime('%d %B %Y')}")
    doc.add_paragraph()

    for line in CONTACT_DETAILS:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    add_heading(doc, "1. Executive Submission Letter", level=1)
    add_body(
        doc,
        "Dear Plumstead Neighbourhood Watch Committee,",
    )
    add_body(
        doc,
        "Gambew Consulting hereby submits this proposal to partner with Plumstead Neighbourhood "
        "Watch on a focused 8-week transformation programme. The objective is to elevate the current "
        "platform from a strong feature-complete implementation to an operationally excellent, "
        "modern, dynamic, and high-value digital community service.",
    )
    add_body(
        doc,
        "This submission is grounded in implementation evidence from the current repository and route-level "
        "workflow review, combined with a practical delivery framework, executive governance controls, and "
        "a clear commercial model aligned with your requested terms.",
    )

    add_heading(doc, "2. Profile Overview: Gareth Bew", level=1)
    add_body(
        doc,
        "Gareth Bew is the Chief Executive Officer of BEW AND CO (PTY) LTD and a senior digital "
        "transformation leader with delivery exposure across South Africa, the Middle East, and broader "
        "international enterprise environments.",
    )
    add_bullets(
        doc,
        [
            "Leadership profile spans retail, banking, and financial services transformation programmes.",
            "Proven track record in data analytics, AI, machine learning, and enterprise modernization.",
            "Senior appointments include Woolworths, Alshaya Group, TFG, Sanlam, Nedbank, and related enterprise portfolios.",
            "Demonstrated achievement in markdown optimization transformation with significant decision-cycle acceleration.",
            "Experience in enterprise-scale value chain, analytics, and operational transformation programmes.",
            "Executive governance exposure through EXCO and steering-level programme leadership.",
            "Academic and professional credentials include MSc Global Business Management and postgraduate digital strategy/marketing studies.",
        ],
    )
    add_body(
        doc,
        "Leadership positioning for PNW: technology-enabled operational modernization with measurable "
        "community outcomes, disciplined delivery governance, and practical transition support.",
    )
    add_body(
        doc,
        "Profile evidence source: LinkedIn-generated profile document (`c:\\Users\\garet\\Downloads\\Profile.pdf`, dated 18 February 2026).",
    )

    add_heading(doc, "3. C-Level Current Platform Assessment", level=1)
    add_body(
        doc,
        "The current PNW platform is materially advanced and should be treated as a modern baseline rather "
        "than a greenfield rebuild. Core digital journeys and data structures are already implemented.",
    )
    add_bullets(
        doc,
        [
            "Implemented service journeys: incidents, events/RSVP, registration, zone finder, volunteer, vacation watch, start-scheme, contact.",
            "Modern architecture in place: Next.js, React, Prisma/Postgres, Clerk auth, typed actions and validation schemas.",
            "Data coverage supports operational expansion: users, incidents, events, engagement records, and enquiry workflows.",
            "Key executive gap: operational assurance maturity (evidence completion, signoff discipline, release hardening).",
            "Governance artefacts exist but require closure for production-grade programme confidence.",
            "Release/runtime consistency requires hardening to reduce deployment variance and execution risk.",
        ],
    )
    add_body(
        doc,
        "Strategic conclusion: transition from feature-complete build status to an operationally excellent, "
        "committee-governed community platform with sustained service quality and accountability.",
    )

    add_heading(doc, "4. 8-Week Transformation Delivery Plan", level=1)
    for title, points in WEEK_PLAN:
        add_heading(doc, title, level=2)
        add_bullets(doc, points)

    add_heading(doc, "5. Deliverables and Governance Model", level=1)
    add_bullets(
        doc,
        [
            "Full platform audit and executive modernization report.",
            "Transformation roadmap with milestone gates and acceptance points.",
            "Committee presentation pack and stakeholder walkthrough materials.",
            "Verification report with evidence matrix completion.",
            "Expansion recommendations for next-phase service value.",
            "Six-month support/maintenance/transition service activation.",
            "Included hosting and technical infrastructure operations.",
        ],
    )
    add_body(
        doc,
        "Governance cadence: weekly checkpoint reporting, milestone reviews, decision log tracking, and formal "
        "committee signoff at each delivery gate.",
    )

    add_heading(doc, "6. Financial Proposition", level=1)
    table = doc.add_table(rows=len(FINANCIAL_TABLE), cols=2)
    table.style = "Table Grid"
    for r_idx, row in enumerate(FINANCIAL_TABLE):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value

    add_heading(doc, "7. Public Interfaces and Operational Data Domains", level=1)
    add_body(doc, "Primary public service interfaces in scope:")
    add_bullets(doc, PUBLIC_INTERFACES)
    add_body(doc, "Core operational data domains in scope:")
    add_bullets(doc, DATA_DOMAINS)
    add_body(
        doc,
        "Governance references: release runbook, regression manifest, evidence matrix, and action log.",
    )

    add_heading(doc, "8. Acceptance, QA, and Credibility Scenarios", level=1)
    add_bullets(
        doc,
        [
            "Financial terms match approved commercial model exactly.",
            "Profile claims align with extracted profile evidence.",
            "Contact and signature details match approved primary details.",
            "Registration flow completes with zone linkage.",
            "Incident reporting persists correctly.",
            "Event RSVP add/remove is reliable.",
            "Contact, volunteer, vacation-watch, and start-scheme workflows validate and persist correctly.",
            "Route smoke tests pass for all core public routes.",
            "Evidence matrix and multi-role signoff controls are closed for release readiness.",
        ],
    )

    add_heading(doc, "9. Assumptions and Defaults", level=1)
    add_bullets(
        doc,
        [
            "Proposal audience: Plumstead Neighbourhood Watch Committee.",
            "Delivery commitment: 8 weeks.",
            "Submission pack: DOCX and PDF.",
            "Pricing format: milestone narrative plus clear commercial table.",
            "Profile section is an executive summary, supported by extracted profile evidence.",
            "Primary contact details are as supplied in the submission request.",
        ],
    )

    add_heading(doc, "10. Closeout and Sign-Off", level=1)
    add_body(
        doc,
        "Proposal validity: 30 days from the submission date.",
    )
    add_body(doc, "")
    add_body(doc, "Approved for Plumstead Neighbourhood Watch:")
    add_body(doc, "Name: ______________________________")
    add_body(doc, "Role: _______________________________")
    add_body(doc, "Signature: __________________________")
    add_body(doc, "Date: _______________________________")
    add_body(doc, "")
    add_body(doc, "Submitted by:")
    add_body(doc, "Gareth Bew")
    add_body(doc, "Gambew Consulting")

    doc.save(DOCX_PATH)


def build_pdf() -> None:
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    styles = getSampleStyleSheet()
    style_title = ParagraphStyle(
        "cover_title",
        parent=styles["Title"],
        fontSize=26,
        leading=30,
        alignment=1,
        spaceAfter=8,
    )
    style_subtitle = ParagraphStyle(
        "cover_subtitle",
        parent=styles["Normal"],
        fontSize=12,
        leading=16,
        alignment=1,
    )
    style_h1 = ParagraphStyle(
        "h1",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        spaceBefore=10,
        spaceAfter=8,
    )
    style_h2 = ParagraphStyle(
        "h2",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceBefore=6,
        spaceAfter=4,
    )
    style_body = ParagraphStyle(
        "body",
        parent=styles["BodyText"],
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    )
    style_bullet = ParagraphStyle(
        "bullet",
        parent=style_body,
        leftIndent=14,
        bulletIndent=4,
    )

    story = []
    story.append(Paragraph("GAMBEW CONSULTING", style_title))
    story.append(Paragraph("Professional Tender Submission", style_subtitle))
    story.append(Spacer(1, 18))
    story.append(
        Paragraph("PNW World-Class Platform Transformation Proposal", style_subtitle)
    )
    story.append(
        Paragraph("Prepared for: Plumstead Neighbourhood Watch Committee", style_subtitle)
    )
    story.append(Paragraph(f"Date: {TODAY.strftime('%d %B %Y')}", style_subtitle))
    story.append(Spacer(1, 16))
    for line in CONTACT_DETAILS:
        story.append(Paragraph(line, style_subtitle))
    story.append(PageBreak())

    story.append(Paragraph("1. Executive Submission Letter", style_h1))
    story.append(Paragraph("Dear Plumstead Neighbourhood Watch Committee,", style_body))
    story.append(
        Paragraph(
            "Gambew Consulting hereby submits this proposal to partner with Plumstead Neighbourhood "
            "Watch on a focused 8-week transformation programme. The objective is to elevate the current "
            "platform from a strong feature-complete implementation to an operationally excellent, "
            "modern, dynamic, and high-value digital community service.",
            style_body,
        )
    )
    story.append(
        Paragraph(
            "This submission is grounded in implementation evidence from the current repository and route-level "
            "workflow review, combined with a practical delivery framework, executive governance controls, and "
            "a clear commercial model aligned with the requested terms.",
            style_body,
        )
    )

    story.append(Paragraph("2. Profile Overview: Gareth Bew", style_h1))
    story.append(
        Paragraph(
            "Gareth Bew is the Chief Executive Officer of BEW AND CO (PTY) LTD and a senior digital "
            "transformation leader with delivery exposure across South Africa, the Middle East, and broader "
            "international enterprise environments.",
            style_body,
        )
    )
    profile_points = [
        "Leadership profile spans retail, banking, and financial services transformation programmes.",
        "Proven track record in data analytics, AI, machine learning, and enterprise modernization.",
        "Senior appointments include Woolworths, Alshaya Group, TFG, Sanlam, Nedbank, and related enterprise portfolios.",
        "Demonstrated achievement in markdown optimization transformation with significant decision-cycle acceleration.",
        "Experience in enterprise-scale value chain, analytics, and operational transformation programmes.",
        "Executive governance exposure through EXCO and steering-level programme leadership.",
        "Academic and professional credentials include MSc Global Business Management and postgraduate digital strategy/marketing studies.",
    ]
    for item in profile_points:
        story.append(Paragraph(item, style_bullet, bulletText="•"))
    story.append(
        Paragraph(
            "Leadership positioning for PNW: technology-enabled operational modernization with measurable "
            "community outcomes, disciplined delivery governance, and practical transition support.",
            style_body,
        )
    )
    story.append(
        Paragraph(
            "Profile evidence source: c:\\Users\\garet\\Downloads\\Profile.pdf (LinkedIn export, 18 February 2026).",
            style_body,
        )
    )

    story.append(Paragraph("3. C-Level Current Platform Assessment", style_h1))
    story.append(
        Paragraph(
            "The current PNW platform is materially advanced and should be treated as a modern baseline rather "
            "than a greenfield rebuild. Core digital journeys and data structures are already implemented.",
            style_body,
        )
    )
    c_level_points = [
        "Implemented service journeys: incidents, events/RSVP, registration, zone finder, volunteer, vacation watch, start-scheme, contact.",
        "Modern architecture in place: Next.js, React, Prisma/Postgres, Clerk auth, typed actions and validation schemas.",
        "Data coverage supports operational expansion: users, incidents, events, engagement records, and enquiry workflows.",
        "Key executive gap: operational assurance maturity (evidence completion, signoff discipline, release hardening).",
        "Governance artefacts exist but require closure for production-grade programme confidence.",
        "Release/runtime consistency requires hardening to reduce deployment variance and execution risk.",
    ]
    for item in c_level_points:
        story.append(Paragraph(item, style_bullet, bulletText="•"))
    story.append(
        Paragraph(
            "Strategic conclusion: transition from feature-complete build status to an operationally excellent, "
            "committee-governed community platform with sustained service quality and accountability.",
            style_body,
        )
    )

    story.append(Paragraph("4. 8-Week Transformation Delivery Plan", style_h1))
    for title, points in WEEK_PLAN:
        story.append(Paragraph(title, style_h2))
        for p in points:
            story.append(Paragraph(p, style_bullet, bulletText="•"))

    story.append(Paragraph("5. Deliverables and Governance Model", style_h1))
    deliverables = [
        "Full platform audit and executive modernization report.",
        "Transformation roadmap with milestone gates and acceptance points.",
        "Committee presentation pack and stakeholder walkthrough materials.",
        "Verification report with evidence matrix completion.",
        "Expansion recommendations for next-phase service value.",
        "Six-month support/maintenance/transition service activation.",
        "Included hosting and technical infrastructure operations.",
    ]
    for d in deliverables:
        story.append(Paragraph(d, style_bullet, bulletText="•"))
    story.append(
        Paragraph(
            "Governance cadence: weekly checkpoint reporting, milestone reviews, decision log tracking, and "
            "formal committee signoff at each delivery gate.",
            style_body,
        )
    )

    story.append(Paragraph("6. Financial Proposition", style_h1))
    table = Table(
        FINANCIAL_TABLE,
        colWidths=[52 * mm, 120 * mm],
        hAlign="LEFT",
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B1F3A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D7E2")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7FAFF")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 8))

    story.append(Paragraph("7. Public Interfaces and Operational Data Domains", style_h1))
    story.append(Paragraph("Primary public service interfaces in scope:", style_body))
    for p in PUBLIC_INTERFACES:
        story.append(Paragraph(p, style_bullet, bulletText="•"))
    story.append(Paragraph("Core operational data domains in scope:", style_body))
    for d in DATA_DOMAINS:
        story.append(Paragraph(d, style_bullet, bulletText="•"))
    story.append(
        Paragraph(
            "Governance references: release runbook, regression manifest, evidence matrix, and action log.",
            style_body,
        )
    )

    story.append(Paragraph("8. Acceptance, QA, and Credibility Scenarios", style_h1))
    qa_points = [
        "Financial terms match approved commercial model exactly.",
        "Profile claims align with extracted profile evidence.",
        "Contact and signature details match approved primary details.",
        "Registration flow completes with zone linkage.",
        "Incident reporting persists correctly.",
        "Event RSVP add/remove is reliable.",
        "Contact, volunteer, vacation-watch, and start-scheme workflows validate and persist correctly.",
        "Route smoke tests pass for all core public routes.",
        "Evidence matrix and multi-role signoff controls are closed for release readiness.",
    ]
    for q in qa_points:
        story.append(Paragraph(q, style_bullet, bulletText="•"))

    story.append(Paragraph("9. Assumptions and Defaults", style_h1))
    assumptions = [
        "Proposal audience: Plumstead Neighbourhood Watch Committee.",
        "Delivery commitment: 8 weeks.",
        "Submission pack: DOCX and PDF.",
        "Pricing format: milestone narrative plus clear commercial table.",
        "Profile section is an executive summary, supported by extracted profile evidence.",
        "Primary contact details are as supplied in the submission request.",
    ]
    for a in assumptions:
        story.append(Paragraph(a, style_bullet, bulletText="•"))

    story.append(Paragraph("10. Closeout and Sign-Off", style_h1))
    story.append(Paragraph("Proposal validity: 30 days from submission date.", style_body))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Approved for Plumstead Neighbourhood Watch:", style_body))
    story.append(Paragraph("Name: ______________________________", style_body))
    story.append(Paragraph("Role: _______________________________", style_body))
    story.append(Paragraph("Signature: __________________________", style_body))
    story.append(Paragraph("Date: _______________________________", style_body))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Submitted by:", style_body))
    story.append(Paragraph("Gareth Bew", style_body))
    story.append(Paragraph("Gambew Consulting", style_body))

    doc.build(story)


def main() -> None:
    ROOT.joinpath("scripts").mkdir(parents=True, exist_ok=True)
    build_docx()
    build_pdf()
    print(f"Generated DOCX: {DOCX_PATH}")
    print(f"Generated PDF:  {PDF_PATH}")


if __name__ == "__main__":
    main()
